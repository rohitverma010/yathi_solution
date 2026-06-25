"""Streamlit app — Telugu video → English SRT converter.

Left column : video/SRT upload  OR  URL input
Right column: translated SRT output + download button
"""
from __future__ import annotations

import io
import os
import re
import subprocess
import tempfile
from pathlib import Path

import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI

load_dotenv(dotenv_path=Path(__file__).parent / ".env.local")
load_dotenv(dotenv_path=Path(__file__).parent / ".env")

CHUNK_SECONDS = 900          # 15-minute chunks (same as pipeline)
SUPPORTED_VIDEO_TYPES = ["mp4", "mov", "avi", "mkv", "webm", "m4v"]
SUPPORTED_TYPES = SUPPORTED_VIDEO_TYPES + ["srt"]

# ── helpers ──────────────────────────────────────────────────────────────────

def _seconds_to_srt_time(s: float) -> str:
    h = int(s // 3600)
    m = int((s % 3600) // 60)
    sec = int(s % 60)
    ms = int(round((s - int(s)) * 1000))
    return f"{h:02d}:{m:02d}:{sec:02d},{ms:03d}"


def segments_to_srt(segments: list[dict]) -> str:
    lines: list[str] = []
    for i, seg in enumerate(segments, 1):
        start = _seconds_to_srt_time(seg["start"])
        end   = _seconds_to_srt_time(seg["end"])
        text  = seg["text"].strip()
        lines.append(f"{i}\n{start} --> {end}\n{text}\n")
    return "\n".join(lines)


def extract_and_chunk_audio(video_path: str, chunk_dir: str) -> list[str]:
    """Use ffmpeg to extract audio and split into CHUNK_SECONDS chunks."""
    out_pattern = str(Path(chunk_dir) / "chunk_%03d.mp3")
    cmd = [
        "ffmpeg", "-y", "-i", video_path,
        "-vn", "-ac", "1", "-ar", "16000", "-b:a", "32k",
        "-f", "segment", "-segment_time", str(CHUNK_SECONDS),
        "-reset_timestamps", "1",
        out_pattern,
    ]
    proc = subprocess.run(cmd, capture_output=True, text=True)
    if proc.returncode != 0:
        raise RuntimeError(f"ffmpeg failed:\n{proc.stderr[-1500:]}")
    chunks = sorted(Path(chunk_dir).glob("chunk_*.mp3"))
    if not chunks:
        raise RuntimeError("ffmpeg produced no audio chunks.")
    return [str(p) for p in chunks]


def translate_chunks(client: OpenAI, chunk_paths: list[str], status_fn) -> list[dict]:
    """Send each chunk to Whisper translations endpoint; return stitched segments."""
    all_segments: list[dict] = []
    for i, path in enumerate(chunk_paths):
        status_fn(f"Translating chunk {i + 1} / {len(chunk_paths)} …")
        offset = i * CHUNK_SECONDS
        with open(path, "rb") as f:
            resp = client.audio.translations.create(
                model="whisper-1",
                file=f,
                response_format="verbose_json",
            )
        for s in (getattr(resp, "segments", None) or []):
            all_segments.append({
                "start": float(s.start) + offset,
                "end":   float(s.end)   + offset,
                "text":  (s.text or "").strip(),
            })
    return all_segments


def process_video(video_bytes: bytes, filename: str, status_fn) -> str:
    """Full pipeline: bytes → SRT string."""
    api_key = os.environ.get("OPENAI_API_KEY", "")
    if not api_key:
        raise RuntimeError(
            "OPENAI_API_KEY not set. Add it to .env.local or your shell environment."
        )
    client = OpenAI(api_key=api_key)

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = str(Path(tmpdir) / filename)
        Path(tmp_path).write_bytes(video_bytes)

        status_fn("Extracting & splitting audio with ffmpeg …")
        chunk_paths = extract_and_chunk_audio(tmp_path, tmpdir)

        segments = translate_chunks(client, chunk_paths, status_fn)

    status_fn("Formatting SRT …")
    return segments_to_srt(segments)


def _parse_srt_blocks(content: str) -> list[dict]:
    blocks = re.split(r"\n\s*\n", content.strip())
    segments = []
    for block in blocks:
        lines = block.strip().splitlines()
        if len(lines) < 3:
            continue
        try:
            idx = int(lines[0].strip())
            times = lines[1].strip()
            text = "\n".join(lines[2:]).strip()
            segments.append({"index": idx, "times": times, "text": text})
        except (ValueError, IndexError):
            continue
    return segments


def translate_srt(srt_content: str, status_fn) -> str:
    """Translate SRT text from Telugu to English using GPT, preserving timestamps."""
    api_key = os.environ.get("OPENAI_API_KEY", "")
    if not api_key:
        raise RuntimeError(
            "OPENAI_API_KEY not set. Add it to .env.local or your shell environment."
        )
    client = OpenAI(api_key=api_key)
    segments = _parse_srt_blocks(srt_content)
    if not segments:
        raise RuntimeError("Could not parse SRT file — check the format.")

    BATCH = 50
    translated: list[str] = []

    for start in range(0, len(segments), BATCH):
        batch = segments[start : start + BATCH]
        status_fn(
            f"Translating subtitles {start + 1}–{min(start + BATCH, len(segments))} / {len(segments)} …"
        )
        source = "\n---\n".join(s["text"] for s in batch)
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Translate the following Telugu subtitle texts to English. "
                        "Return ONLY the translations separated by '---', one per entry, "
                        "preserving order and count exactly. No extra commentary."
                    ),
                },
                {"role": "user", "content": source},
            ],
        )
        parts = resp.choices[0].message.content.strip().split("---")
        parts = [p.strip() for p in parts]
        while len(parts) < len(batch):
            parts.append("")
        translated.extend(parts[: len(batch)])

    lines: list[str] = []
    for i, seg in enumerate(segments):
        lines.append(f"{seg['index']}\n{seg['times']}\n{translated[i]}\n")
    return "\n".join(lines)


def srt_to_docx(srt_content: str) -> bytes:
    """Convert SRT text to a .docx file; returns raw bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading("Translated Subtitles", level=1)
    blocks = re.split(r"\n\s*\n", srt_content.strip())
    for block in blocks:
        lines = block.strip().splitlines()
        if len(lines) < 3:
            continue
        timecode = lines[1].strip()
        text = " ".join(lines[2:]).strip()
        p = doc.add_paragraph()
        run_time = p.add_run(f"{timecode}\n")
        run_time.font.size = Pt(8)
        run_time.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run_text = p.add_run(text)
        run_text.font.size = Pt(11)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def download_from_url(url: str, dest_dir: str, status_fn) -> str:
    """Download video from a URL using yt-dlp; return local file path."""
    status_fn("Downloading video from URL …")
    out_template = str(Path(dest_dir) / "downloaded.%(ext)s")
    cmd = [
        "yt-dlp",
        "-f", "bestvideo[ext=mp4]+bestaudio[ext=m4a]/best[ext=mp4]/best",
        "--merge-output-format", "mp4",
        "-o", out_template,
        url,
    ]
    proc = subprocess.run(cmd, capture_output=True, text=True)
    if proc.returncode != 0:
        raise RuntimeError(f"yt-dlp failed:\n{proc.stderr[-1500:]}")
    matches = list(Path(dest_dir).glob("downloaded.*"))
    if not matches:
        raise RuntimeError("yt-dlp produced no output file.")
    return str(matches[0])


# ── UI ───────────────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Telugu → English Subtitle Converter",
    layout="wide",
)

st.title("Telugu → English Subtitle Converter")
st.caption("Upload a Telugu video or SRT, or paste a URL. Audio is translated to English and formatted as an SRT subtitle file.")

left, right = st.columns(2)

# ── LEFT: upload / URL ────────────────────────────────────────────────────────
with left:
    st.subheader("Input")
    tab_file, tab_url = st.tabs(["Upload File", "From URL"])

    uploaded = None
    url_input = ""

    with tab_file:
        uploaded = st.file_uploader(
            "Choose a video or SRT file",
            type=SUPPORTED_TYPES,
            help=f"Supported formats: {', '.join(SUPPORTED_TYPES).upper()}",
        )
        if uploaded:
            if uploaded.name.lower().endswith(".srt"):
                st.info(f"**SRT file:** {uploaded.name}  \n**Size:** {uploaded.size / 1000:.1f} KB")
            else:
                st.video(uploaded)
                st.info(
                    f"**File:** {uploaded.name}  \n"
                    f"**Size:** {uploaded.size / 1_000_000:.1f} MB"
                )

    with tab_url:
        url_input = st.text_input(
            "Video URL",
            placeholder="https://www.youtube.com/watch?v=...",
            help="Paste any URL supported by yt-dlp (YouTube, Vimeo, etc.)",
        )
        if url_input:
            st.info(f"Will download: `{url_input}`")

    has_input = bool(uploaded or url_input.strip())
    convert_btn = st.button(
        "Convert to English SRT",
        type="primary",
        use_container_width=True,
        disabled=not has_input,
    )

# ── RIGHT: output ─────────────────────────────────────────────────────────────
with right:
    st.subheader("English SRT Output")

    if "srt_content" not in st.session_state:
        st.session_state["srt_content"] = ""
    if "srt_filename" not in st.session_state:
        st.session_state["srt_filename"] = "output.srt"

    if convert_btn and has_input:
        status_box = st.empty()

        def update_status(msg: str) -> None:
            status_box.info(msg)

        with st.spinner("Processing — this may take a few minutes …"):
            try:
                # ── SRT upload ────────────────────────────────────────────
                if uploaded and uploaded.name.lower().endswith(".srt"):
                    srt_raw = uploaded.getvalue().decode("utf-8", errors="replace")
                    srt_text = translate_srt(srt_raw, update_status)
                    stem = Path(uploaded.name).stem
                    st.session_state["srt_filename"] = f"{stem}.en.srt"

                # ── video upload ──────────────────────────────────────────
                elif uploaded:
                    video_bytes = uploaded.getvalue()
                    srt_text = process_video(video_bytes, uploaded.name, update_status)
                    stem = Path(uploaded.name).stem
                    st.session_state["srt_filename"] = f"{stem}.en.srt"

                # ── URL ───────────────────────────────────────────────────
                else:
                    with tempfile.TemporaryDirectory() as tmpdir:
                        video_path = download_from_url(url_input.strip(), tmpdir, update_status)
                        video_bytes = Path(video_path).read_bytes()
                        srt_text = process_video(video_bytes, Path(video_path).name, update_status)
                    st.session_state["srt_filename"] = "downloaded.en.srt"

                st.session_state["srt_content"] = srt_text
                status_box.success("Translation complete!")
            except Exception as exc:
                status_box.error(f"Error: {exc}")

    if st.session_state["srt_content"]:
        st.text_area(
            "SRT content",
            value=st.session_state["srt_content"],
            height=480,
            label_visibility="collapsed",
        )
        dl_left, dl_right = st.columns(2)
        with dl_left:
            st.download_button(
                label="Download .srt",
                data=st.session_state["srt_content"].encode("utf-8"),
                file_name=st.session_state["srt_filename"],
                mime="text/plain",
                use_container_width=True,
            )
        with dl_right:
            docx_stem = Path(st.session_state["srt_filename"]).stem
            st.download_button(
                label="Download .docx",
                data=srt_to_docx(st.session_state["srt_content"]),
                file_name=f"{docx_stem}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
    else:
        st.info("Upload a file or paste a URL on the left, then click **Convert**.")
